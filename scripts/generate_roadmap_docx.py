#!/usr/bin/env python3
"""Generate arbgrid-roadmap.docx — an owner-facing remaining-work snapshot.

Reads authoritative state from the live codebase (config.STRATEGY_LAYERS, scans/,
tests/, dashboard_ui.py, HANDOFF.json, git log) and writes a Word document at
the repo root that consolidates everything that is still outstanding.

Usage:
    python scripts/generate_roadmap_docx.py
    python scripts/generate_roadmap_docx.py --output /custom/path/roadmap.docx

The .docx is treated as a generated artifact — it is regenerable, gitignored,
and the script is the source of truth for its structure.
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


LAYER_NAMES = {
    1: "Layer 1 — Pure Arbitrage (Risk-Free)",
    2: "Layer 2 — Near-Arbitrage (Near Risk-Free)",
    3: "Layer 3 — Market Making & Liquidity Provision (Low Risk)",
    4: "Layer 4 — Informed / Statistical Edge (Moderate Risk)",
    5: "Layer 5 — Capital & Execution Optimization (Multiplier)",
}

# Map opportunity-type names from STRATEGY_LAYERS to the scan module they live in.
# Names that don't appear here fall back to a snake_case guess.
OPP_TO_SCAN = {
    "Binary": "binary",
    "NegRisk": "negrisk",
    "Cross": "cross",
    "MultiCross": "multi_cross",
    "TriangularCross": "triangular",
    "KalshiBinary": "kalshi",
    "KalshiMulti": "kalshi",
    "GeminiBinary": "gemini",
    "GeminiMulti": "gemini",
    "IBKRBinary": "ibkr",
    "BetfairBackAll": "betfair",
    "BetfairBackLay": "betfair",
    "SmarketsBackAll": "smarkets",
    "SmarketsBackLay": "smarkets",
    "SXBetBackAll": "sxbet",
    "SXBetBackLay": "sxbet",
    "MatchbookBackAll": "matchbook",
    "MatchbookBackLay": "matchbook",
    "Spread": "spread",
    "BracketArb": "bracket",
    "ConditionalArb": "conditional",
    "NWayArb": None,  # no scan module yet
    "ResolutionSnipeOpp": "resolution",
    "StalePriceOpp": "stale",
    "FeePromo": "fee_promo",
    "SettlementTimingArb": "settlement_timing",
    "NewMarketMispricing": "new_market",
    "APIOutageArb": "api_outage",
    "MarketMake": None,  # market_maker.py, not in scans/
    "CrossPlatformMM": "cross_mm",
    "VolatilityAdjustedMM": None,  # not yet built
    "LeadLagMM": None,
    "ToxicFlowPause": None,
    "EventDivergence": None,  # event_monitor.py
    "ConvergenceOpp": "convergence",
    "SocialSentiment": "social_sentiment",
    "ExpertDivergence": "expert_divergence",
    "CalibratedSignal": None,  # calibration_tracker.py
    "InsiderPattern": "insider_pattern",
    "CrossCategoryCorrelation": "cross_category",
    "OpportunityCostScore": None,  # capital_optimizer.py
    "MarginOptimization": None,
    "TaxAwarePosition": None,
    "WithdrawalTiming": None,
    # Additional types likely present
    "PolymarketRewards": "rewards",
    "KalshiRewards": "rewards",
    "Imbalance": "imbalance",
    "LogicalArb": "logical_arb",
    "TimeDecay": "time_decay",
    "WhaleCopy": "whale_copy",
    "Correlated": "correlated",
}

# Non-scan home modules (some strategies live in dedicated modules, not scans/)
NON_SCAN_HOMES = {
    "MarketMake": "market_maker.py",
    "EventDivergence": "event_monitor.py",
    "CalibratedSignal": "calibration_tracker.py",
    "OpportunityCostScore": "capital_optimizer.py",
    "MarginOptimization": "capital_optimizer.py",
    "TaxAwarePosition": "capital_optimizer.py",
    "WithdrawalTiming": "capital_optimizer.py",
}


def _build_module_test_map() -> dict[str, set[str]]:
    """Grep tests/*.py for `from scans.<x>` / `import scans.<x>` / `from <home>`
    patterns and return a {module_stem: {test_filename, ...}} mapping.

    This is the test-resolution mechanism — Codex flagged that the previous
    one-file-per-strategy heuristic missed `test_new_strategies.py` and other
    grouped test files.
    """
    tests_dir = REPO_ROOT / "tests"
    if not tests_dir.exists():
        return {}
    pat = re.compile(
        r"(?:from\s+scans\.(\w+)|import\s+scans\.(\w+)|from\s+(\w+)\s+import)"
    )
    module_to_tests: dict[str, set[str]] = {}
    for test_file in tests_dir.glob("test_*.py"):
        try:
            text = test_file.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue
        for m in pat.finditer(text):
            module = m.group(1) or m.group(2) or m.group(3)
            if module:
                module_to_tests.setdefault(module, set()).add(test_file.name)
    return module_to_tests


# Semantic-flag patterns the script greps for inside each home module to
# populate the Notes column. Captures the case Codex flagged: STUB / dead /
# quarantined cannot be inferred from filesystem presence alone.
# Curated notes for strategies whose semantic state is not detectable from a
# simple source-code grep — e.g. quarantine guards in config.py, by-design
# ceilings, or strategies that intentionally don't have a scans/ module.
CURATED_NOTES = {
    "SXBetBackAll": "SX Bet quarantined in config.validate_config — place_order() sends unsigned JSON; DRY_RUN-only until EIP-712 signing lands",
    "SXBetBackLay": "SX Bet quarantined in config.validate_config — place_order() sends unsigned JSON; DRY_RUN-only until EIP-712 signing lands",
    "NWayArb": "Layer 1 — no scan module yet; needs implementation",
    "LeadLagMM": "Layer 3 — feature flag exists, no scan/MM logic yet",
    "ToxicFlowPause": "Layer 3 — feature flag exists, no implementation yet",
    "VolatilityAdjustedMM": "Layer 3 — feature flag exists, no implementation yet",
    "OpportunityCostScore": "INFRA by design — scoring class in capital_optimizer.py, not a standalone opp-producing scan",
    "MarginOptimization": "INFRA by design — cross-platform collateral router in capital_optimizer.py",
    "TaxAwarePosition": "INFRA by design — loss harvesting / gain deferral in capital_optimizer.py",
    "WithdrawalTiming": "INFRA by design — withdrawal-delay model in capital_optimizer.py",
    "MarketMake": "INFRA — lives in market_maker.py (QuoteEngine, InventoryTracker, QuoteManager)",
    "EventDivergence": "INFRA — lives in event_monitor.py + signal_aggregator.py",
    "CalibratedSignal": "INFRA — lives in calibration_tracker.py",
}

SEMANTIC_FLAGS = [
    (re.compile(r"\bTODO\b", re.I), "TODO"),
    (re.compile(r"\bFIXME\b", re.I), "FIXME"),
    (re.compile(r"raise\s+NotImplementedError", re.I), "NotImplementedError"),
    (re.compile(r"\bquarantin", re.I), "quarantined"),
    (re.compile(r"\bstub\b", re.I), "stub"),
    (re.compile(r"by[-\s]?design\s+ceiling", re.I), "by-design ceiling"),
    (re.compile(r"\bdead\s+code\b", re.I), "dead code"),
    (re.compile(r"\bread[-\s]?only\b", re.I), "read-only"),
]


def _scan_semantic_flags(path: Path) -> str:
    if not path.exists():
        return ""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return ""
    hits: list[str] = []
    for pattern, label in SEMANTIC_FLAGS:
        count = len(pattern.findall(text))
        if count > 0:
            hits.append(f"{label}×{count}" if count > 1 else label)
    return "; ".join(hits)


def collect_strategy_status() -> list[dict]:
    """Walk config.STRATEGY_LAYERS and resolve scan + test + home for each.

    Uses grep-based test resolution (per Codex round-1 feedback) so grouped
    test files like `test_new_strategies.py` are counted correctly.
    """
    from config import STRATEGY_LAYERS

    scans_dir = REPO_ROOT / "scans"
    module_to_tests = _build_module_test_map()

    rows = []
    for opp_type, layer in sorted(STRATEGY_LAYERS.items(), key=lambda kv: (kv[1], kv[0])):
        scan_name = OPP_TO_SCAN.get(opp_type, _snake(opp_type))
        notes_source: Path | None = None

        if scan_name is None:
            scan_module = NON_SCAN_HOMES.get(opp_type, "—")
            scan_path = REPO_ROOT / scan_module if scan_module != "—" else None
            scan_exists = scan_path is not None and scan_path.exists()
            scan_label = scan_module if scan_exists else (f"({scan_module})" if scan_module != "—" else "—")
            notes_source = scan_path if scan_exists else None
            home_stems = [scan_module.removesuffix(".py")] if scan_module != "—" else []
        else:
            scan_path = scans_dir / f"{scan_name}.py"
            scan_exists = scan_path.exists()
            scan_label = f"scans/{scan_name}.py" if scan_exists else f"(scans/{scan_name}.py missing)"
            notes_source = scan_path if scan_exists else None
            home_stems = [scan_name]

        # Test resolution: scan_name match OR home_stem match in the grep map
        test_files: set[str] = set()
        if scan_name and scan_name in module_to_tests:
            test_files |= module_to_tests[scan_name]
        for stem in home_stems:
            if stem in module_to_tests:
                test_files |= module_to_tests[stem]
        test_exists = bool(test_files)

        scanned = _scan_semantic_flags(notes_source) if notes_source else ""
        curated = CURATED_NOTES.get(opp_type, "")
        notes = " | ".join(n for n in (curated, scanned) if n)

        if scan_exists and test_exists:
            auto_status = "BUILT"
        elif scan_exists and not test_exists:
            auto_status = "PARTIAL (no test imports module)"
        elif not scan_exists and opp_type in NON_SCAN_HOMES:
            auto_status = "INFRA (lives outside scans/)" if (REPO_ROOT / NON_SCAN_HOMES[opp_type]).exists() else "MISSING"
        else:
            auto_status = "MISSING"

        # Per Codex round-2 M3: curated override beats filesystem-derived status
        status = CURATED_STATUS_OVERRIDE.get(opp_type, auto_status)
        v2_status = V2_STATUS.get(opp_type, "—")

        rows.append(
            {
                "opp_type": opp_type,
                "layer": layer,
                "module": scan_label,
                "test": ", ".join(sorted(test_files)) if test_files else "—",
                "status": status,
                "auto_status": auto_status,
                "v2_status": v2_status,
                "flag_default": get_flag_default(opp_type),
                "notes": notes,
            }
        )
    return rows


def _snake(name: str) -> str:
    s1 = re.sub("(.)([A-Z][a-z]+)", r"\1_\2", name)
    return re.sub("([a-z0-9])([A-Z])", r"\1_\2", s1).lower()


def count_innerhtml() -> int:
    path = REPO_ROOT / "dashboard_ui.py"
    if not path.exists():
        return 0
    return path.read_text(encoding="utf-8", errors="ignore").count("innerHTML")


def dashboard_test_exists() -> bool:
    return (REPO_ROOT / "tests" / "test_dashboard_ui.py").exists()


def sxbet_signing_status() -> str:
    path = REPO_ROOT / "sxbet_api.py"
    if not path.exists():
        return "sxbet_api.py not found"
    text = path.read_text(encoding="utf-8", errors="ignore")
    has_eip712 = "EIP-712" in text or "eip712" in text or "_typed_data" in text
    sends_unsigned = "unsigned" in text.lower() and "place_order" in text
    if has_eip712 and not sends_unsigned:
        return "EIP-712 signing implemented (verify with live order)"
    if sends_unsigned:
        return "place_order() still sends unsigned JSON — read-only quarantine"
    return "needs inspection (no EIP-712 markers, no explicit unsigned warning)"


def read_handoff() -> dict:
    path = REPO_ROOT / ".planning" / "HANDOFF.json"
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


# Per Codex round-2 finding M1: never publish HANDOFF.json verbatim; allowlist
# fields and run a redaction pass for common secret-shaped patterns. Anything
# matched is replaced with [REDACTED]. The owner sees the placeholder and can
# look at the original .planning/HANDOFF.json if they need the raw value.
HANDOFF_FIELD_ALLOWLIST_TOP = {"version", "timestamp", "phase", "phase_name", "status"}
HANDOFF_FIELD_ALLOWLIST_BLOCKER = {"description", "type", "workaround"}

SECRET_PATTERNS = [
    re.compile(r"\bsk_[A-Za-z0-9_\-]{12,}", re.I),
    re.compile(r"\bpk_[A-Za-z0-9_\-]{12,}", re.I),
    re.compile(r"\bbearer\s+[A-Za-z0-9._\-]+", re.I),
    re.compile(r"\beyJ[A-Za-z0-9._\-]{20,}"),  # JWT-like
    re.compile(r"https?://[^/\s:@]+:[^@\s]+@\S+", re.I),  # URL with creds
    re.compile(r"\b[A-Z_]+_KEY\s*[=:]\s*[A-Za-z0-9._\-]+", re.I),
    re.compile(r"\b0x[a-fA-F0-9]{40,}\b"),  # ETH addresses / private keys
]


def _redact(text: str) -> str:
    for pat in SECRET_PATTERNS:
        text = pat.sub("[REDACTED]", text)
    return text


STALE_HANDOFF_DAYS = 30


def safe_handoff_blocker(handoff: dict) -> tuple[str, str, int | None]:
    """Return (timestamp, redacted_blocker_text_or_stale_placeholder, days_old).

    Per Codex round-3 H1: when the file is older than STALE_HANDOFF_DAYS, drop
    the blocker text entirely. A stale "100% rejection" claim is worse than no
    claim — it focuses owner attention on a potentially-resolved issue.
    """
    timestamp = handoff.get("timestamp", "unknown")
    age = _handoff_age_days(timestamp)
    blockers = handoff.get("blockers", []) or []
    if not blockers:
        return timestamp, "No active blocker recorded.", age

    if age is not None and age > STALE_HANDOFF_DAYS:
        return (
            timestamp,
            (
                f"Blocker text suppressed — HANDOFF.json is {age} days old "
                f"(threshold: {STALE_HANDOFF_DAYS} days). Reverify against live "
                "state before treating any pre-existing blocker as current. Run "
                "the Pre-Action Live-State Verification Checklist (next section) "
                "to determine the actual current state."
            ),
            age,
        )

    b = blockers[0]
    filtered = {k: v for k, v in b.items() if k in HANDOFF_FIELD_ALLOWLIST_BLOCKER}
    parts = []
    for key in ("description", "type", "workaround"):
        if key in filtered:
            parts.append(f"{key}: {filtered[key]}")
    body = "\n".join(parts) if parts else "Blocker present but no allowlisted fields."
    return timestamp, _redact(body), age


def _handoff_age_days(timestamp: str) -> int | None:
    try:
        ts = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
    except (ValueError, AttributeError):
        return None
    delta = datetime.now(ts.tzinfo) - ts
    return delta.days


# Per Codex round-3 M1: surface flag default so BUILT-status entries are visibly
# tagged as "code present but disabled by default" when applicable.
OPP_TO_FLAG = {
    "Binary": "BINARY_ENABLED",
    "NegRisk": "NEGRISK_ENABLED",
    "KalshiBinary": "KALSHI_ENABLED",
    "KalshiMulti": "KALSHI_MULTI_ENABLED",
    "GeminiBinary": "GEMINI_ENABLED",
    "GeminiMulti": "GEMINI_ENABLED",
    "IBKRBinary": "IBKR_ENABLED",
    "Cross": "CROSS_ENABLED",
    "MultiCross": "MULTI_CROSS_ENABLED",
    "TriangularCross": "TRIANGULAR_ENABLED",
    "Spread": "SPREAD_ENABLED",
    "BracketArb": "BRACKET_ARB_ENABLED",
    "ConditionalArb": "CONDITIONAL_ARB_ENABLED",
    "NWayArb": "NWAY_ARB_ENABLED",
    "SXBetBackAll": "SXBET_ENABLED",
    "SXBetBackLay": "SXBET_ENABLED",
    "BetfairBackAll": "BETFAIR_ENABLED",
    "BetfairBackLay": "BETFAIR_ENABLED",
    "SmarketsBackAll": "SMARKETS_ENABLED",
    "SmarketsBackLay": "SMARKETS_ENABLED",
    "MatchbookBackAll": "MATCHBOOK_ENABLED",
    "MatchbookBackLay": "MATCHBOOK_ENABLED",
    "ResolutionSnipeOpp": "RESOLUTION_ENABLED",
    "StalePriceOpp": "STALE_ENABLED",
    "FeePromo": "FEE_PROMO_ENABLED",
    "SettlementTimingArb": "SETTLEMENT_TIMING_ENABLED",
    "NewMarketMispricing": "NEW_MARKET_MISPRICING_ENABLED",
    "APIOutageArb": "API_OUTAGE_ARB_ENABLED",
    "MarketMake": "MM_ENABLED",
    "CrossPlatformMM": "CROSS_MM_ENABLED",
    "VolatilityAdjustedMM": "MM_VOLATILITY_ADJUSTED_ENABLED",
    "LeadLagMM": "LEAD_LAG_MM_ENABLED",
    "ToxicFlowPause": "MM_TOXIC_FLOW_ENABLED",
    "EventDivergence": "EVENT_MONITOR_ENABLED",
    "ConvergenceOpp": "CONVERGENCE_ENABLED",
    "SocialSentiment": "SOCIAL_SENTIMENT_ENABLED",
    "ExpertDivergence": "EXPERT_DIVERGENCE_ENABLED",
    "CalibratedSignal": "CALIBRATION_WEIGHTING_ENABLED",
    "InsiderPattern": "INSIDER_PATTERN_ENABLED",
    "CrossCategoryCorrelation": "CROSS_CATEGORY_ENABLED",
    "OpportunityCostScore": "OPPORTUNITY_COST_SCORING_ENABLED",
    "MarginOptimization": "MARGIN_EFFICIENCY_ENABLED",
    "TaxAwarePosition": "TAX_AWARE_ENABLED",
    "WithdrawalTiming": "WITHDRAWAL_TIMING_ENABLED",
}


def get_flag_default(opp_type: str) -> str:
    """Return 'true' | 'false' | '—' for the feature flag controlling opp_type."""
    flag = OPP_TO_FLAG.get(opp_type)
    if not flag:
        return "—"
    import config  # type: ignore

    val = getattr(config, flag, None)
    if val is None:
        return "—"
    return "true" if val else "false"


def find_unregistered_opp_types() -> list[tuple[str, str]]:
    """Grep scans/*.py for `type="X"` literals; return (opp_type, scan_file)
    tuples where opp_type is not in config.STRATEGY_LAYERS. This catches the
    registration gap where scan code emits an opp type the layer-mapping
    has never been updated to include.
    """
    from config import STRATEGY_LAYERS

    pat = re.compile(r'type\s*[=:]\s*["\'](\w+)["\']')
    scans_dir = REPO_ROOT / "scans"
    seen: dict[str, str] = {}
    if not scans_dir.exists():
        return []
    for path in scans_dir.glob("*.py"):
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue
        for m in pat.finditer(text):
            opp = m.group(1)
            # Heuristic: opp types start with a capital letter
            if opp[:1].isupper() and opp not in STRATEGY_LAYERS:
                seen.setdefault(opp, f"scans/{path.name}")
    return sorted(seen.items())


def git_provenance() -> dict:
    """Return {head, branch, dirty, script_sha256} for the cover page."""
    info = {"head": "unknown", "branch": "unknown", "dirty": False, "script_sha256": ""}
    try:
        info["head"] = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"], cwd=REPO_ROOT, text=True
        ).strip()
        info["branch"] = subprocess.check_output(
            ["git", "rev-parse", "--abbrev-ref", "HEAD"], cwd=REPO_ROOT, text=True
        ).strip()
        status = subprocess.check_output(
            ["git", "status", "--porcelain"], cwd=REPO_ROOT, text=True
        )
        info["dirty"] = bool(status.strip())
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass
    try:
        import hashlib

        script_path = Path(__file__)
        info["script_sha256"] = hashlib.sha256(script_path.read_bytes()).hexdigest()[:12]
    except OSError:
        pass
    return info


def recent_commits(n: int = 15) -> list[str]:
    try:
        out = subprocess.check_output(
            ["git", "log", "--oneline", f"-{n}"], cwd=REPO_ROOT, text=True
        )
        return [line.strip() for line in out.strip().splitlines() if line.strip()]
    except subprocess.CalledProcessError:
        return []


# Per Codex round-2 finding M3: status auto-computed from filesystem + grep can
# silently promote known-PARTIAL strategies to BUILT (e.g., #28 WhaleCopy has a
# scan AND a test, but docs/strategy-framework-v2.md flagged it PARTIAL for a
# real reason). These overrides take precedence over the auto-computed status.
CURATED_STATUS_OVERRIDE = {
    # Layer 1 — SX Bet exchange execution still blocked
    "SXBetBackAll": "PARTIAL (place_order quarantined)",
    "SXBetBackLay": "PARTIAL (place_order quarantined)",
    # Layer 4 — original v2 doc flagged these PARTIAL/STUB pre-refiner work;
    # refiners landed in PRs #12, #14, #15 but Codex correctly notes the
    # framework doc has not been re-blessed. Defer the BUILT label until the
    # framework doc is reconciled.
    "TimeDecay": "BUILT (was PARTIAL in v2; refiner landed PR #12)",
    "WhaleCopy": "PARTIAL (was STUB in v2; decoder + refiner landed PRs #14, decoder; verify against framework before promoting)",
    "Correlated": "PARTIAL (was STUB in v2; refiner landed PR #15; verify before promoting)",
}


# Per Codex round-2 finding M3: surface the historical v2 framework status so
# the matrix shows what the 2026-05-09 reconciliation said for #1-29. Strategies
# #30-#49 have no v2 row.
V2_STATUS = {
    # Layer 1
    "Binary": "BUILT (#1)",
    "NegRisk": "BUILT (#2)",
    "KalshiBinary": "BUILT (#1)",
    "KalshiMulti": "BUILT (#2)",
    "GeminiBinary": "BUILT (#1)",
    "GeminiMulti": "BUILT (#2)",
    "IBKRBinary": "BUILT (#1)",
    "Cross": "BUILT (#3)",
    "MultiCross": "BUILT (#4)",
    "TriangularCross": "BUILT (#5)",
    "BetfairBackAll": "PARTIAL (#6)",
    "BetfairBackLay": "PARTIAL (#6)",
    "SmarketsBackAll": "PARTIAL (#6)",
    "SmarketsBackLay": "PARTIAL (#6)",
    "SXBetBackAll": "PARTIAL (#6 — place_order broken)",
    "SXBetBackLay": "PARTIAL (#6 — place_order broken)",
    "MatchbookBackAll": "PARTIAL (#6)",
    "MatchbookBackLay": "PARTIAL (#6)",
    "Spread": "BUILT (#21)",
    # Layer 2
    "ResolutionSnipeOpp": "BUILT (#7)",
    "StalePriceOpp": "BUILT (#8)",
    "FeePromo": "BUILT (#9, PR #10)",
    # Layer 3
    "MarketMake": "BUILT (#10)",
    "CrossPlatformMM": "BUILT (#11, PR #10)",
    "PolymarketRewards": "BUILT (#22)",
    "KalshiRewards": "BUILT (#23)",
    # Layer 4
    "EventDivergence": "BUILT (#13)",
    "ConvergenceOpp": "BUILT (#14)",
    "Imbalance": "BUILT (#24)",
    "LogicalArb": "BUILT (#25)",
    "TimeDecay": "PARTIAL (#26 — dead Stage 2 refiner)",
    "WhaleCopy": "PARTIAL (#28 — calldata parser MVP stub)",
    "Correlated": "STUB (#29 — TODO marker only)",
    # Layer 5
}


# ---------------------------------------------------------------------------
# Word doc rendering
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)


def add_callout(doc: Document, title: str, body: str, color: str = "FFE599") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, color)
    p1 = cell.paragraphs[0]
    run = p1.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    for r in p2.runs:
        r.font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def build_document(output: Path) -> None:
    strategy_rows = collect_strategy_status()
    handoff = read_handoff()
    innerhtml = count_innerhtml()
    has_dash_test = dashboard_test_exists()
    sxbet = sxbet_signing_status()
    commits = recent_commits(15)

    by_layer: dict[int, list[dict]] = {1: [], 2: [], 3: [], 4: [], 5: []}
    for row in strategy_rows:
        by_layer[row["layer"]].append(row)

    counts = {
        "total": len(strategy_rows),
        "built": sum(1 for r in strategy_rows if r["status"] == "BUILT"),
        "partial": sum(1 for r in strategy_rows if r["status"].startswith("PARTIAL")),
        "infra": sum(1 for r in strategy_rows if r["status"].startswith("INFRA")),
        "missing": sum(1 for r in strategy_rows if r["status"] == "MISSING"),
    }

    doc = Document()
    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -------- Cover --------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("arbgrid — Remaining-Work Roadmap")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "Owner-facing alignment reference — automated, regenerable from live codebase"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}    |    "
        f"Owner: Jonathon Tamm    |    Repo: arbgrid (formerly polymarket-arb-scanner)"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Provenance line per Codex round-2 M4
    prov = git_provenance()
    prov_para = doc.add_paragraph()
    prov_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dirty_label = " (DIRTY — uncommitted changes)" if prov["dirty"] else " (clean)"
    prov_run = prov_para.add_run(
        f"git: {prov['branch']} @ {prov['head']}{dirty_label}    |    "
        f"generator sha256: {prov['script_sha256']}"
    )
    prov_run.font.size = Pt(9)
    prov_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    add_callout(
        doc,
        "How to read this document",
        "This .docx is generated by scripts/generate_roadmap_docx.py. It reads "
        "config.STRATEGY_LAYERS, the scans/ and tests/ directories, dashboard_ui.py, "
        ".planning/HANDOFF.json, and git log to compute the current state. The narrative "
        "(blockers, sequencing, recommendations) is hand-curated. Rerun the script any "
        "time to refresh strategy counts and recent-commit lists. The .docx is gitignored; "
        "the script is the source of truth.",
        color="DEEBF7",
    )

    doc.add_page_break()

    # -------- 1. Mission & Success --------
    add_heading(doc, "Mission & Success Criteria", level=1)
    doc.add_paragraph(
        "arbgrid is a 24/7 automated prediction-market trading bot deployed on Railway. "
        "It scans 8 trading platforms plus 2 read-only signal sources for opportunities "
        "across 5 risk layers. Detection, revalidation, execution, risk management, "
        "market making, monitoring, and backtesting are all in-repo."
    )
    doc.add_paragraph("Success criteria (from CLAUDE.md):")
    for crit in [
        "Net positive P&L in trades.db over a 7-day live trading period",
        "<5% false-positive rate on detected opportunities (manual spot-check vs platforms)",
        "At least one profitable autonomous round-trip trade with no human intervention",
    ]:
        doc.add_paragraph(crit, style="List Bullet")

    # -------- History (Codex round-1 ask: preserve the original narrative) --------
    add_heading(doc, "History — How We Got Here", level=1)
    doc.add_paragraph(
        "The arbgrid project did not start at 43 opportunity types. It started as a "
        "polymarket-only binary-arb scanner, grew into an 8-platform / 20-strategy bot for "
        "v1.0, was reconciled into a 29-strategy / 5-layer framework in v2 "
        "(docs/strategy-framework-v2.md, 2026-05-09), and then expanded again in PR #22 to "
        "add 20 new strategies (#30-#49). This section preserves the v2 framing so the "
        "current matrix is read as evolution, not contradiction."
    )

    add_heading(doc, "Original v2 status (29 strategies, snapshot 2026-05-09)", level=2)
    add_table(
        doc,
        ["Status", "Count", "Strategies"],
        [
            ["BUILT", "22", "1, 2, 3, 4, 5, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 19, 21, 22, 23, 24, 25"],
            ["PARTIAL", "6", "6 (SX Bet quarantine), 18 (by-design ceiling), 20, 26, 27, 28"],
            ["STUB", "1", "29 (correlated pairs — TODO only)"],
            ["NOT BUILT", "0", "—"],
        ],
    )

    add_heading(doc, "5-phase remediation roadmap and what closed each phase", level=2)
    phase_rows = [
        [
            "Phase 1 — Quick wins (1 week)",
            "SX Bet quarantine, resolution-window env var, dashboard host/password hardening",
            "Closed by PR #18 (2026-05-09 era)",
        ],
        [
            "Phase 2 — Finish Layer 4 incomplete scans (#26, #27, #28, #29)",
            "Stage 2 refiners + whale calldata decoder + correlation tracker",
            "Closed by PRs #12, #14, #15, plus whale_copy_decoder.py and correlation_tracker.py",
        ],
        [
            "Phase 3 — Close original-framework gap (#20 tuning loop)",
            "scripts/tune.py — backtest-driven threshold tuner",
            "Closed by PR #16 (scripts/tune.py)",
        ],
        [
            "Phase 4 — Hardening",
            "Rate-limiter audit + dedicated test files for missing modules + dashboard XSS",
            "Partially closed by PR #17 (rate-limiter + 4 test files). Dashboard XSS still open.",
        ],
        [
            "Phase 5 — Future-state (platform-gated)",
            "Extend #18 auto-rebalance corridor to other platforms",
            "Open — gated on external platform APIs",
        ],
    ]
    add_table(doc, ["Phase", "Scope", "Disposition"], phase_rows)

    add_heading(doc, "PR #22 — expansion to 49 strategies", level=2)
    doc.add_paragraph(
        "After Phases 1-4 closed, the codebase added 20 net-new strategies (#30-#49) covering "
        "conditional arb, bracket arb, n-way arb, settlement-timing arb, new-market mispricing, "
        "API-outage arb, three new MM variants, social sentiment, expert divergence, calibrated "
        "signals, insider patterns, cross-category correlation, and four capital-optimization "
        "modules (opportunity cost, margin, tax-aware, withdrawal timing). Infrastructure plus a "
        "subset of strategy implementations landed in one PR. Some #30-#49 entries are still "
        "MISSING or live as INFRA-only (capital optimizer classes without a scan module) — "
        "see the current status matrix below."
    )

    # -------- Snapshot --------
    add_heading(doc, "Current Status Snapshot", level=1)
    doc.add_paragraph(
        f"Total opportunity types registered in config.STRATEGY_LAYERS: "
        f"{counts['total']}    |    BUILT: {counts['built']}    |    "
        f"PARTIAL: {counts['partial']}    |    INFRA-only: {counts['infra']}    |    "
        f"MISSING: {counts['missing']}"
    )

    add_heading(doc, "Layer breakdown", level=2)
    layer_rows = []
    for layer in (1, 2, 3, 4, 5):
        entries = by_layer.get(layer, [])
        built = sum(1 for e in entries if e["status"] == "BUILT")
        partial = sum(1 for e in entries if e["status"].startswith("PARTIAL"))
        infra = sum(1 for e in entries if e["status"].startswith("INFRA"))
        missing = sum(1 for e in entries if e["status"] == "MISSING")
        layer_rows.append(
            [
                LAYER_NAMES[layer].split(" — ")[0],
                LAYER_NAMES[layer].split(" — ")[1],
                str(len(entries)),
                str(built),
                str(partial),
                str(infra),
                str(missing),
            ]
        )
    add_table(
        doc,
        ["Layer", "Description", "Total", "BUILT", "PARTIAL", "INFRA", "MISSING"],
        layer_rows,
    )

    add_heading(doc, "Platform inventory", level=2)
    platform_rows = [
        ["Polymarket", "Ethereum CLOB", "Buy + Sell", "BUILT"],
        ["Kalshi", "RSA-PSS headers", "Buy + Sell", "BUILT"],
        ["Betfair", "SSO + API key", "Back + Lay", "BUILT"],
        ["Smarkets", "API key session", "Back + Lay", "BUILT"],
        ["SX Bet", "API key session", "Quarantined", "READ-ONLY (EIP-712 pending)"],
        ["Matchbook", "Username / password", "Back + Lay (0% commission)", "BUILT"],
        ["Gemini Predictions", "HMAC-SHA384", "Buy + Sell (1% / 5%)", "BUILT"],
        ["IBKR ForecastEx", "TWS API (IB Gateway)", "BUY-only, LMT, $0 comm.", "BUILT"],
        ["Metaculus", "Public REST", "Read-only signal", "BUILT"],
        ["Manifold Markets", "Public REST", "Read-only signal", "BUILT"],
    ]
    add_table(doc, ["Platform", "Auth", "Trade direction", "Status"], platform_rows)

    # -------- Critical blocker --------
    # Per Codex round-2 M1+M2: redact secrets + flag staleness
    add_heading(doc, "Critical Operational Blocker — Profitability Validation", level=1)
    handoff_ts, blocker_text, age_days = safe_handoff_blocker(handoff)

    if age_days is not None and age_days > 30:
        add_callout(
            doc,
            f"⚠ STALE — HANDOFF.json was last updated {age_days} days ago",
            "This blocker snapshot is from "
            f"{handoff_ts}. Verify against live trades.db, Railway logs, and current "
            "calibration reports before treating it as the current state. The blocker "
            "may already be resolved.",
            color="FFC7CE",
        )

    add_callout(
        doc,
        f"HANDOFF.json blocker (timestamp: {handoff_ts}, allowlisted fields, secrets redacted)",
        blocker_text,
        color="F8CBAD",
    )
    doc.add_paragraph(
        "Recommended Sprint 0 action: query the most recent 7 days of trades.db. If detection "
        "rate is non-zero and execution rate is still 0, the revalidation gate in "
        "executor.py:_revalidate* is still rejecting everything and needs to be re-tuned. If "
        "execution rate is non-zero, the blocker is stale — close HANDOFF.json and update "
        "STATE.md before any further strategy work."
    )
    doc.add_paragraph("Files to inspect:")
    for f in [
        "executor.py — _revalidate_binary / _revalidate_negrisk / _revalidate_cross / etc.",
        "trades.db — query trades and opportunities tables",
        ".planning/calibration-reports/ — latest entries for live calibration state",
        "Railway logs — current 24h scan/execution counters",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    # -------- Strategy matrix --------
    add_heading(doc, "Strategy Status Matrix — All Registered Opportunity Types", level=1)
    doc.add_paragraph(
        "Generated programmatically from config.STRATEGY_LAYERS, scans/, and tests/. "
        "Test coverage is detected by grep-searching every tests/*.py for "
        "`from scans.<name>` / `import scans.<name>` / `from <home_module>` patterns, "
        "so grouped tests like test_new_strategies.py are counted correctly."
    )
    doc.add_paragraph(
        "Status legend: BUILT (code present + test imports module — does NOT certify "
        "production readiness or that the feature flag is enabled); PARTIAL (module "
        "exists, no test imports it); INFRA (lives outside scans/, e.g. market_maker.py "
        "/ capital_optimizer.py); MISSING (feature flag registered, no implementation). "
        "Flag default shows the *_ENABLED env var default in config.py — a BUILT "
        "strategy with flag default=false is shipped but disabled until explicitly "
        "enabled. Notes column surfaces semantic signals (quarantines, by-design "
        "ceilings, INFRA placement)."
    )
    for layer in (1, 2, 3, 4, 5):
        entries = by_layer.get(layer, [])
        if not entries:
            continue
        add_heading(doc, LAYER_NAMES[layer], level=2)
        rows = [
            [
                e["opp_type"],
                e["module"],
                e["test"],
                e["status"],
                e["flag_default"],
                e["v2_status"],
                e["notes"] or "—",
            ]
            for e in entries
        ]
        add_table(
            doc,
            [
                "Opportunity type",
                "Module",
                "Test file(s)",
                "Current status",
                "Flag default",
                "v2 status (2026-05-09)",
                "Notes",
            ],
            rows,
        )

    # -------- Strategy gaps --------
    add_heading(doc, "Strategy Gaps — What's Not Yet Built", level=1)

    unregistered = find_unregistered_opp_types()
    if unregistered:
        add_callout(
            doc,
            f"⚠ Registration gap — {len(unregistered)} opp types in code are missing from config.STRATEGY_LAYERS",
            "These types are emitted by scan code (`type=\"...\"` literals) but never "
            "appear in the layer-mapping. Layer-aware revalidation, position sizing, and "
            "P&L attribution all key off STRATEGY_LAYERS, so unregistered opp types fall "
            "through to defaults and may be under-protected. Fix: add each to "
            "config.STRATEGY_LAYERS with its correct layer.\n\n"
            + "\n".join(f"• {opp} (in {src})" for opp, src in unregistered),
            color="FFC7CE",
        )

    missing_rows = [r for r in strategy_rows if r["status"] == "MISSING"]
    partial_rows = [r for r in strategy_rows if r["status"].startswith("PARTIAL")]

    if missing_rows:
        doc.add_paragraph(
            f"{len(missing_rows)} strategies have feature flags but no implementation:"
        )
        for r in missing_rows:
            doc.add_paragraph(
                f"{r['opp_type']} (Layer {r['layer']}) — {r['module']}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No registered strategies are entirely missing implementations.")

    if partial_rows:
        doc.add_paragraph(
            f"{len(partial_rows)} strategies are PARTIAL (module exists, dedicated test missing):"
        )
        for r in partial_rows:
            doc.add_paragraph(
                f"{r['opp_type']} (Layer {r['layer']}) — {r['module']}",
                style="List Bullet",
            )

    # -------- 6. Hardening --------
    add_heading(doc, "Remaining Hardening Work", level=1)
    hardening_rows = [
        [
            "Dashboard XSS",
            f"{innerhtml} occurrences of `innerHTML` in dashboard_ui.py",
            "Replace with textContent or DOM API; medium severity",
            "Open" if innerhtml > 0 else "Closed",
        ],
        [
            "Dashboard UI test coverage",
            "tests/test_dashboard_ui.py " + ("exists" if has_dash_test else "does NOT exist"),
            "Add coverage for HTML render + endpoints",
            "Closed" if has_dash_test else "Open",
        ],
        [
            "SX Bet EIP-712 signing",
            sxbet,
            "Implement EIP-712 typed-data signing for place_order; unblocks live exec",
            "Open" if "unsigned" in sxbet.lower() or "inspection" in sxbet.lower() else "Closed",
        ],
        [
            "Rate-limiter audit (Betfair/Smarkets)",
            "PR #17 claims complete — spot-verify all external calls route through rate_limiter",
            "Sample 5-10 methods per platform; confirm decorator presence",
            "Verify",
        ],
    ]
    add_table(doc, ["Item", "Current state", "Fix", "Status"], hardening_rows)

    # -------- 7. Platform-gated --------
    add_heading(doc, "Platform-Gated Future Work", level=1)
    doc.add_paragraph(
        "These items can't be built until external platforms expose required APIs. Track "
        "for awareness, but do not sequence into sprints until the gate opens."
    )
    for item in [
        "Extend #18 auto-rebalance corridor beyond Gemini ↔ Polymarket. Currently no other "
        "platform (Kalshi, Betfair, Smarkets, SX Bet, Matchbook, IBKR) exposes programmatic "
        "withdraw/deposit endpoints. Gated on platform behavior.",
        "SX Bet EIP-712 signing — also listed under Hardening because it is partially in "
        "our control. If SX Bet ever ships an official signing library or schema reference, "
        "the implementation cost drops significantly.",
        "STRAT-08 Betfair/Smarkets in-play scalping — requires live-event data feed; "
        "currently out of v2.0 scope per REQUIREMENTS.md.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # -------- 8. Operational path --------
    add_heading(doc, "Operational Path to Profitability", level=1)
    doc.add_paragraph(
        "These are not engineering tasks — they are owner-action items required to validate "
        "the system end-to-end. Pulled from REQUIREMENTS.md (LIVE-01 through LIVE-06)."
    )
    for item in [
        "Capital funding — confirm balance on all 8 trading platforms; document minimums.",
        "Live trading window — run the bot for a continuous 7-day period with execution enabled.",
        "P&L validation — query trades.db post-window for net P&L, win rate, FP rate.",
        "False-positive audit — sample 20-50 detected opportunities, manually verify against "
        "platform UIs; confirm <5% FP rate.",
        "Round-trip verification — confirm at least one fully autonomous profitable trade "
        "(entry, exit, P&L > 0).",
        "Credential rotation — verify credential_health.py is firing alerts for any platform "
        "approaching credential expiry.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # -------- Pre-Action Live-State Verification Checklist (Codex round-3 M4) --------
    add_heading(doc, "Pre-Action Live-State Verification Checklist", level=1)
    doc.add_paragraph(
        "This document is repo-local — it does not query live production state. Before "
        "acting on ANY sprint, ANY blocker, or ANY status claim, run this checklist to "
        "establish actual current state. Stale repo metadata + stale handoff data + a "
        "drift between code presence and prod readiness can otherwise drive bad decisions."
    )
    checklist = [
        ("Deployed revision",  "railway status; railway logs --tail 20 | head -5 (compare to local git HEAD)"),
        ("Execution rate (last 24h)", "python -c \"import sqlite3; c=sqlite3.connect('trades.db').cursor(); print(c.execute(\\\"SELECT COUNT(*) FROM opportunities WHERE created_at > datetime('now','-1 day')\\\").fetchone(), c.execute(\\\"SELECT COUNT(*) FROM trades WHERE created_at > datetime('now','-1 day')\\\").fetchone())\""),
        ("Per-strategy P&L (last 7d)", "python scripts/pnl_report.py --days 7"),
        ("Calibration trend", "ls -t .planning/calibration-reports/ | head -5; cat $(ls -t .planning/calibration-reports/ | head -1)"),
        ("Credential health", "python -m credential_health (or check Railway /healthz output)"),
        ("Active feature flags in prod", "railway run env | grep -E '_ENABLED|MM_|FEE_' (or check Railway dashboard Variables)"),
        ("Dashboard health endpoints", "curl https://<railway-domain>/healthz; curl https://<railway-domain>/metrics | head"),
        ("Open alerts / loss streaks", "Check webhook destination (Slack/Discord) for last 24h; review alerting.py state"),
        ("Platform quota headroom", "Check Polymarket / Kalshi / Betfair / Smarkets API quota usage in their respective dashboards"),
    ]
    add_table(doc, ["Probe", "Command"], [[name, cmd] for name, cmd in checklist])
    doc.add_paragraph(
        "If any probe fails or reveals a divergence from this document, treat THIS document "
        "as wrong and act on the live data."
    )

    # -------- Sprint sequence (Codex round-3 M2 + M3: deps + rollback) --------
    add_heading(doc, "Recommended Execution Sequence", level=1)
    doc.add_paragraph(
        "Sprints are ordered so monitoring/hardening precedes strategy expansion (this "
        "preserves the original v2.0 phase dependency: Deploy → Monitor → Liquidity → "
        "Signal Strategies → Structural Alpha). Each sprint has explicit Dependencies "
        "and Rollback / containment lines so an incident response has a known kill-switch."
    )

    sprints = [
        (
            "Sprint 0 — Reverify current state",
            "Run the Pre-Action Live-State Verification Checklist (previous section). "
            "Document findings. If the HANDOFF.json blocker (100% rejection, 2026-04-01) "
            "is still active, scope a fix. If resolved, update STATE.md and HANDOFF.json.",
            "None — this IS the first action.",
            "N/A — checklist is read-only.",
            "Live-state probes complete; written summary of which blockers are real today.",
        ),
        (
            "Sprint 1 — Close monitoring / hardening gaps",
            f"Dashboard XSS: replace {innerhtml} innerHTML occurrences with textContent. Add "
            "tests/test_dashboard_ui.py. Verify credential_health.py is firing on schedule. "
            "Verify alerting.py loss-streak + zero-opportunity detectors are wired. Verify "
            "WS heartbeat detection (HARD-01) is active.",
            "Sprint 0.",
            "Dashboard XSS fix is purely additive (textContent never executes); revert via "
            "single commit revert. Alert/heartbeat wiring is config — disable the flag.",
            "0 innerHTML occurrences; test_dashboard_ui.py passes; alerting webhook fires on "
            "test-induced loss streak.",
        ),
        (
            "Sprint 2 — Build MISSING strategy scans",
            f"Implement the {len(missing_rows)} MISSING entries (NWayArb, LeadLagMM, "
            "ToxicFlowPause, VolatilityAdjustedMM). Each ships behind a default-OFF flag. "
            "Each gets a dedicated test file that imports the module.",
            "Sprint 1 (monitoring must catch new-strategy regressions).",
            "Each new strategy is gated by its *_ENABLED flag — flip to false at the Railway "
            "env var level for instant disable. No code rollback needed.",
            "All MISSING entries either implemented (flag default=false) OR explicitly "
            "removed from STRATEGY_LAYERS with rationale.",
        ),
        (
            "Sprint 3 — Register unregistered opp types + add dedicated tests",
            f"Add the unregistered opp types (Correlated, Imbalance, NewsSnipe, TimeDecay, "
            "WhaleCopy) to config.STRATEGY_LAYERS so layer-aware revalidation and "
            "position-sizing apply. Backfill dedicated tests for the PARTIAL entries.",
            "Sprint 1.",
            "Registration is a config change — revert the STRATEGY_LAYERS edit in one commit. "
            "Tests are additive.",
            "0 unregistered opp types reported by this generator; PARTIAL count drops.",
        ),
        (
            "Sprint 4 — SX Bet EIP-712 signing",
            "Implement EIP-712 typed-data signing in sxbet_api.py:place_order(). Test on SX "
            "Bet testnet (NOT mainnet). Remove the validate_config() quarantine block. Add "
            "tests/test_sxbet_eip712.py.",
            "Sprints 1 + 3.",
            "validate_config() block stays in place until the testnet round-trip succeeds. "
            "If a mainnet order misbehaves, set DRY_RUN=true and SXBET_ENABLED=false via "
            "Railway env vars to halt all SX Bet activity within 1 deploy cycle.",
            "Testnet round-trip (place + cancel) succeeds; PR diff includes EIP-712 unit tests.",
        ),
        (
            "Sprint 5 — Revalidation tuning if Sprint 0 confirmed blocker is live",
            "If Sprint 0 confirmed the 100% rejection rate is still live: investigate "
            "executor.py:_revalidate* methods. Likely causes: WS price cache TTL too tight; "
            "stale_prices guard too aggressive; layer-floor thresholds wrong post-fee-overhaul. "
            "Make ONE change at a time; measure delta over a 24h dry-run window.",
            "Sprint 0 (must confirm blocker is live, not stale).",
            "Each revalidation tuning change goes behind a temporary env var (e.g., "
            "REVAL_TOLERANCE_OVERRIDE). Revert via env var unset, no redeploy needed.",
            "24h dry-run shows 5-30% of detected opportunities passing revalidation (per "
            "EXEC-01 acceptance in REQUIREMENTS.md).",
        ),
        (
            "Sprint 6 — 7-day live profitability window",
            "Fund all 8 trading platforms (LIVE-01..LIVE-06). Enable execution. Run "
            "continuously for 7 days. Produce validation report via "
            "scripts/validation_report.py.",
            "Sprints 0, 1, 5. SX Bet (Sprint 4) optional — can run with SX Bet quarantined.",
            "Kill switch: DRY_RUN=true via Railway env var halts all live trading "
            "immediately. Per-strategy kill: flip *_ENABLED=false. Daily loss limit "
            "(DAILY_LOSS_LIMIT) trips automatic halt if breached.",
            "Net positive P&L; <5% FP rate (per CLAUDE.md success criteria); ≥1 profitable "
            "autonomous round-trip recorded in trades.db.",
        ),
    ]
    for title, body, deps, rollback, acceptance in sprints:
        add_heading(doc, title, level=2)
        doc.add_paragraph(body)
        p = doc.add_paragraph()
        p.add_run("Dependencies: ").bold = True
        p.add_run(deps)
        p = doc.add_paragraph()
        p.add_run("Rollback / containment: ").bold = True
        p.add_run(rollback)
        p = doc.add_paragraph()
        p.add_run("Acceptance: ").bold = True
        p.add_run(acceptance)

    # -------- 10. Recent velocity --------
    add_heading(doc, "Recent Velocity (last 15 commits)", level=1)
    if commits:
        for c in commits:
            doc.add_paragraph(c, style="List Bullet")
    else:
        doc.add_paragraph("Could not read git log.")

    # -------- 11. Source pointers --------
    add_heading(doc, "Authoritative Source Documents", level=1)
    doc.add_paragraph(
        "This .docx is a derived snapshot. The following remain authoritative going forward:"
    )
    for src in [
        "CLAUDE.md — project scope, success criteria, platform inventory",
        "config.py — STRATEGY_LAYERS and all *_ENABLED feature flags",
        "scans/ — actual scan module implementations (one file per scan)",
        "docs/strategy-framework-v2.md — historical strategy taxonomy (1-29 only; treat as context, not status)",
        ".planning/ROADMAP.md — GSD v2.0 phase structure",
        ".planning/REQUIREMENTS.md — EXEC/MON/HARD/STRAT requirement IDs",
        ".planning/STATE.md — current phase pointer and decision log",
        ".planning/HANDOFF.json — last paused-task context (verify freshness)",
        "git log — actual landed work (latest is always the most accurate)",
    ]:
        doc.add_paragraph(src, style="List Bullet")

    doc.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output",
        default=str(REPO_ROOT / "arbgrid-roadmap.docx"),
        help="Output path for the .docx file",
    )
    args = parser.parse_args()
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    build_document(output)
    print(f"Wrote {output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
